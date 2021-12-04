from xlrd import open_workbook
from docx import Document

excel_file = "BEES Active Publications List.xlsx"

wb = open_workbook(filename=excel_file)
ws = wb.sheet_by_index(0)

# store data
manuscript_numbers = []
titles = []
authors = []
abstracts = []
acknowledgements = []
publication_status = []
journals = []
date = []
volume = []
doi = []


def get_all_in_col(row_num, ws, arr):
    first_element = True
    for i in range(ws.nrows):
        if first_element:
            first_element = False
            pass
        else:
            arr.append(ws.cell_value(i, row_num))


'''
To add a column, first add a data storage above
Then add an elif statement like:
elif ws.cell_value(0, i) == "Column_Title":
    get_all_in_col(i, ws, Name_of_Data_Storage)
'''
for i in range(ws.ncols):
    if ws.cell_value(0, i) == "Title":
        get_all_in_col(i, ws, titles)  # Get all titles
    elif ws.cell_value(0, i) == "Authors":
        get_all_in_col(i, ws, authors)  # Get all authors
    elif ws.cell_value(0, i) == "Journal":
        get_all_in_col(i, ws, journals)  # Get all journal names
    elif ws.cell_value(0, i) == "Status":
        get_all_in_col(i, ws, publication_status)  # Get publication status
    elif ws.cell_value(0, i) == "M#":
        get_all_in_col(i, ws, manuscript_numbers)  # Get Bees manuscript number
    elif ws.cell_value(0, i) == "Abstract":
        get_all_in_col(i, ws, abstracts)  # Get Bees manuscript number
    elif ws.cell_value(0, i) == "Acknowledgment":
        get_all_in_col(i, ws, acknowledgements)  # Get Bees manuscript number
    elif ws.cell_value(0, i) == "Date":
        get_all_in_col(i, ws, date)  # Get Bees manuscript number
    elif ws.cell_value(0, i) == "Volume":
        get_all_in_col(i, ws, volume)  # Get Bees manuscript number
    elif ws.cell_value(0, i) == "DOI":
        get_all_in_col(i, ws, doi)  # Get Bees manuscript number

for x in range(len(authors)):
    string = authors[x]
    authors[x] = string.replace(';', ',')  # Names must be separated by commas

for x in range(len(manuscript_numbers)):
    num = manuscript_numbers[x]
    string = str(num)
    if int(num) < 10:
        manuscript_numbers[x] = '00' + string.replace('.0', '')  # Gets rid of decimal
    elif int(num) < 100:
        manuscript_numbers[x] = '0' + string.replace('.0', '')  # Gets rid of decimal
    else:
        manuscript_numbers[x] = string.replace('.0', '')  # Gets rid of decimal

for x in range(len(date)):
    string = str(date[x])
    date[x] = string.replace('.0', '')


# Format of reference: Authors, Title, Journal, Status
# references = open('References.txt', 'wb')
#
# while len(titles) > 0:
#     ref_outstr = authors.pop(0) + ', "' + titles.pop(0) + '", ' + journals.pop(0) + ', ' + publication_status.pop(
#         0) + '\n'
#     ref_outstr = ref_outstr.replace(', ,', ',')  # Get rid of empty elements
#     references.write(ref_outstr.encode('utf-8'))

# Creates BEES Appendix as requested by customer
# appendix = open('Appendix.txt', 'wb')


def get_from_lists(arr):
    if arr:
        return str(arr.pop(0))
    else:
        return "Unknown"


def pop_all():
    get_from_lists(manuscript_numbers)
    titles.pop(0)
    authors.pop(0)
    get_from_lists(abstracts)
    get_from_lists(acknowledgements)
    publication_status.pop(0)
    journals.pop(0)
    get_from_lists(date)
    get_from_lists(volume)
    get_from_lists(doi)


'''
while len(titles) > 0:
    print(len(titles), len(publication_status))
    manuscript_outstr = "BEES Manuscript Number: " + get_from_lists(manuscript_numbers) + "\n\n"
    title_outstr = "Title: " + titles.pop(0) + "\n"
    author_outstr = "Authors: " + authors.pop(0) + "\n\n"
    abstract_outstr = "Abstract: " + get_from_lists(abstracts) + "\n\n"
    ack_outstr = "Acknowledgement: " + get_from_lists(acknowledgements) + "\n\n"
    status_outstr = "Status: " + publication_status.pop(0) + "\n\n"
    journals_outstr = "Journal: " + journals.pop(0) + "\n\n"
    date_outstr = "Date: " + get_from_lists(date) + "\n\n"
    vol_outstr = "Volume: " + get_from_lists(volume) + "\n\n"
    doi_outstr = "DOI: " + get_from_lists(doi) + "\n\n"
    appendix_outstr = (manuscript_outstr + title_outstr + author_outstr + abstract_outstr + ack_outstr + status_outstr +
                       journals_outstr + date_outstr + vol_outstr + doi_outstr + "\n\n\n")
    try:
        if int(status_outstr[8]) <= 5:
            appendix.write(appendix_outstr.encode('utf-8'))
    except ValueError:
        pass
'''

doc = Document()
while len(titles) > 0:
    status = publication_status[0]
    try:
        if len(status) > 0 and int(status[0]) <= 5:
            p = doc.add_paragraph()
            p.add_run('BEES Manuscript Number: ').bold = True
            p.add_run(get_from_lists(manuscript_numbers))

            p = doc.add_paragraph()
            p.add_run('Title: ').bold = True
            p.add_run(titles.pop(0))

            p = doc.add_paragraph()
            p.add_run('Authors: ').bold = True
            p.add_run(authors.pop(0))

            p = doc.add_paragraph()
            p.add_run('Abstract: ').bold = True
            p.add_run(get_from_lists(abstracts))

            p = doc.add_paragraph()
            p.add_run('Acknowledgement: ').bold = True
            p.add_run(get_from_lists(acknowledgements))

            p = doc.add_paragraph()
            p.add_run('Status: ').bold = True
            p.add_run(publication_status.pop(0))

            p = doc.add_paragraph()
            p.add_run('Journal: ').bold = True
            p.add_run(journals.pop(0))

            p = doc.add_paragraph()
            p.add_run('Date: ').bold = True
            p.add_run(get_from_lists(date))

            p = doc.add_paragraph()
            p.add_run('Volume: ').bold = True
            p.add_run(get_from_lists(volume))

            p = doc.add_paragraph()
            p.add_run('DOI: ').bold = True
            p.add_run(get_from_lists(doi))

            doc.add_page_break()

        else:
            pop_all()
    except ValueError:
        pop_all()

doc.save('Appendix.docx')
